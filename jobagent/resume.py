"""Resume handling: extract text, score it against a job description, rewrite it.

When an OpenAI key is configured the analysis and rewrite are LLM-powered. With
no key, a keyword-overlap heuristic still gives a useful match score (rewrite
needs the LLM and will say so).
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from .config import Settings

# Very small stopword list for the heuristic keyword extractor.
_STOPWORDS = set(
    """a an the and or of to in for with on at by from as is are be this that you your we our will
    work job role team company candidate experience years requirements responsibilities including
    etc using ability strong excellent good plus must should have has had they their it its
    about across into per via able who which what when where how all any new use used""".split()
)


def extract_text(path: str | Path) -> str:
    """Extract plain text from a PDF, DOCX, or TXT/MD resume."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume file not found: {p}")
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix in {".docx"}:
        return _extract_docx(p)
    if suffix in {".txt", ".md"}:
        return p.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume type: {suffix} (use PDF, DOCX, TXT, or MD)")


def _extract_pdf(p: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts).strip()


def _extract_docx(p: Path) -> str:
    import docx

    document = docx.Document(str(p))
    return "\n".join(par.text for par in document.paragraphs).strip()


def _keywords(text: str, top: int = 40) -> list[str]:
    words = re.findall(r"[a-zA-Z][a-zA-Z+#.\-]{2,}", text.lower())
    freq: dict[str, int] = {}
    for w in words:
        w = w.strip(".-")
        if len(w) < 3 or w in _STOPWORDS:
            continue
        freq[w] = freq.get(w, 0) + 1
    ranked = sorted(freq, key=lambda k: freq[k], reverse=True)
    return ranked[:top]


def analyze(resume_text: str, jd_text: str, settings: Settings) -> dict[str, Any]:
    """Return a match report. Tries the LLM first, falls back to heuristics."""
    if settings.llm_enabled:
        report = _analyze_llm(resume_text, jd_text, settings)
        if report:
            report["engine"] = "ai"
            return report
    report = _analyze_heuristic(resume_text, jd_text)
    report["engine"] = "keyword"
    return report


def _analyze_heuristic(resume_text: str, jd_text: str) -> dict[str, Any]:
    jd_kw = _keywords(jd_text)
    resume_low = resume_text.lower()
    matched = [k for k in jd_kw if k in resume_low]
    missing = [k for k in jd_kw if k not in resume_low]
    score = round(100 * len(matched) / max(1, len(jd_kw)))
    if score >= 70:
        chance, verdict = "High", "Strong match. You have a good shot at a callback."
    elif score >= 45:
        chance, verdict = "Medium", "Partial match. Add the missing keywords to improve your odds."
    else:
        chance, verdict = "Low", "Weak match. Tailor your resume to this role before applying."
    return {
        "overall_score": score,
        "call_chance": chance,
        "summary": verdict,
        "matched_keywords": matched[:25],
        "missing_keywords": missing[:25],
        "strengths": [f"Mentions: {', '.join(matched[:8])}"] if matched else [],
        "improvements": (
            [f"Add or emphasise: {', '.join(missing[:10])}"] if missing else []
        ),
    }


_ANALYZE_PROMPT = """\
You are an expert technical recruiter and ATS (applicant tracking system). \
Compare the candidate's RESUME against the JOB DESCRIPTION and assess the \
likelihood of getting an interview callback.

Return STRICT JSON with exactly these keys:
{
  "overall_score": <integer 0-100>,
  "call_chance": "High" | "Medium" | "Low",
  "summary": "<2-3 sentence honest assessment>",
  "matched_keywords": ["..."],
  "missing_keywords": ["important skills/keywords from the JD missing in the resume"],
  "strengths": ["specific strengths relative to this role"],
  "improvements": ["concrete, actionable changes to raise the callback chance"]
}
Be honest and specific. Base keywords on the actual JD. Return ONLY the JSON.
"""


def _analyze_llm(resume_text: str, jd_text: str, settings: Settings) -> dict[str, Any] | None:
    try:
        from openai import OpenAI

        client = OpenAI(api_key=settings.openai_api_key)
        resp = client.chat.completions.create(
            model=settings.openai_model,
            response_format={"type": "json_object"},
            temperature=0.2,
            messages=[
                {"role": "system", "content": _ANALYZE_PROMPT},
                {
                    "role": "user",
                    "content": f"JOB DESCRIPTION:\n{jd_text}\n\nRESUME:\n{resume_text}",
                },
            ],
        )
        data = json.loads(resp.choices[0].message.content or "{}")
        data.setdefault("matched_keywords", [])
        data.setdefault("missing_keywords", [])
        data.setdefault("strengths", [])
        data.setdefault("improvements", [])
        data["overall_score"] = int(data.get("overall_score", 0))
        return data
    except Exception:
        return None


_REWRITE_PROMPT = """\
You are an expert resume writer. Rewrite the candidate's resume so it is tailored \
to the target job description and optimised to pass ATS screening, WITHOUT \
inventing facts, employers, degrees, or dates that aren't in the original resume. \
You may rephrase, reorder, surface relevant skills, add a tailored professional \
summary, and naturally incorporate important keywords from the job description \
that the candidate genuinely appears to satisfy.

Output the rewritten resume in clean Markdown with clear sections (Summary, \
Skills, Experience, Education). No commentary before or after -- just the resume.
"""


def rewrite(resume_text: str, jd_text: str, settings: Settings) -> str:
    """Return an improved, tailored resume in Markdown. Requires an LLM key."""
    if not settings.llm_enabled:
        raise RuntimeError(
            "Rewriting needs an OpenAI API key. Add one in Setup to enable AI resume fixing."
        )
    from openai import OpenAI

    client = OpenAI(api_key=settings.openai_api_key)
    resp = client.chat.completions.create(
        model=settings.openai_model,
        temperature=0.4,
        messages=[
            {"role": "system", "content": _REWRITE_PROMPT},
            {
                "role": "user",
                "content": f"JOB DESCRIPTION:\n{jd_text}\n\nCURRENT RESUME:\n{resume_text}",
            },
        ],
    )
    return (resp.choices[0].message.content or "").strip()


def save_docx(markdown_text: str, path: str | Path) -> Path:
    """Save a (lightly-formatted) Markdown resume to a .docx file."""
    import docx

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    document = docx.Document()
    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if not line:
            document.add_paragraph("")
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            document.add_paragraph(line.lstrip()[2:], style="List Bullet")
        else:
            document.add_paragraph(line.replace("**", ""))
    document.save(str(out))
    return out
